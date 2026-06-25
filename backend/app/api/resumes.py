import uuid
import json
import re
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from io import BytesIO

from app.core.database import get_db
from app.api.auth import get_current_user
from app.schemas.resume import (
    ResumeCreate, ResumeUpdate, ResumeResponse, ResumeListResponse,
    ExperienceCreate, ExperienceResponse,
    EducationCreate, EducationResponse,
    SkillCreate, SkillResponse,
    CertificationCreate, CertificationResponse,
    ProjectCreate, ProjectResponse,
)
from app.services import resume_service

router = APIRouter(prefix="/resumes", tags=["resumes"])


def extract_text_from_pdf(file_bytes: bytes) -> str | None:
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
        return text
    except ImportError:
        return None
    except Exception as e:
        raise Exception(f"PDF parsing failed: {str(e)}")


def llm_parse_resume_text(text: str) -> dict | None:
    try:
        import ollama as ollama_sync
        from app.core.config import settings

        prompt = f"""You are an expert resume parser. Convert the following resume text into a structured JSON object.

Resume Text:
{text[:8000]}

Return ONLY a valid JSON object with this exact structure (no markdown, no explanation):
{{
  "name": "Full Name",
  "email": "email@example.com",
  "phone": "phone number",
  "location": "city, state",
  "linkedin": "linkedin URL or empty string",
  "website": "website URL or empty string",
  "summary": "professional summary text",
  "skills": [{{"name": "Skill Name", "proficiency": "intermediate", "category": "language|framework|database|tool|other"}}],
  "experience": [{{"company": "Company Name", "position": "Job Title", "startDate": "YYYY-MM", "endDate": "YYYY-MM or Present", "description": "job description", "highlights": ["highlight 1", "highlight 2"]}}],
  "education": [{{"institution": "School Name", "degree": "Degree", "field": "Field of Study", "startDate": "YYYY-MM", "endDate": "YYYY-MM", "gpa": "GPA or empty string"}}],
  "certifications": [{{"name": "Cert Name", "issuer": "Issuer"}}],
  "projects": [{{"name": "Project Name", "description": "project description", "technologies": ["tech1", "tech2"], "link": "URL or empty string"}}]
}}

Rules:
- Extract ALL information you can find
- Use empty string "" for missing fields
- Use null only for missing array items (return empty array [])
- Dates should be YYYY-MM format when possible
- Categories: language (Python, Java, etc), framework (React, Django, etc), database (MySQL, MongoDB, etc), tool (Docker, AWS, etc)
- Return ONLY the JSON, nothing else"""

        try:
            model = settings.DEFAULT_MODEL
            response = ollama_sync.chat(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                options={"num_gpu": 99, "temperature": 0.1}
            )
            content = response["message"]["content"]

            content = content.strip()
            if content.startswith("```json"):
                content = content[7:]
            elif content.startswith("```"):
                content = content[3:]
            if content.endswith("```"):
                content = content[:-3]
            content = content.strip()

            import json as json_mod
            parsed = json_mod.loads(content)

            return {
                "title": f"{parsed.get('name', '')} Resume" if parsed.get('name') else "Uploaded Resume",
                "template": "modern",
                "summary": parsed.get("summary", ""),
                "target_role": "",
                "personal": {
                    "name": parsed.get("name", ""),
                    "email": parsed.get("email", ""),
                    "phone": parsed.get("phone", ""),
                    "location": parsed.get("location", ""),
                    "website": parsed.get("website", ""),
                    "linkedin": parsed.get("linkedin", ""),
                },
                "experience": parsed.get("experience", []),
                "education": parsed.get("education", []),
                "skills": parsed.get("skills", []),
                "certifications": parsed.get("certifications", []),
                "projects": parsed.get("projects", []),
            }
        except Exception as e:
            print(f"LLM parsing failed, falling back to regex: {e}")
            return None
    except ImportError:
        print("ollama not installed, falling back to regex parsing")
        return None


def extract_text_from_docx(file_bytes: bytes) -> str | None:
    try:
        from docx import Document
        import io
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except ImportError:
        return None
    except Exception as e:
        raise Exception(f"DOCX parsing failed: {str(e)}")


def extract_text_from_tex(content: str) -> str:
    text = content
    text = re.sub(r'\\documentclass.*?\n', '', text, flags=re.DOTALL)
    text = re.sub(r'\\usepackage.*?\n', '', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{document\}', '', text)
    text = re.sub(r'\\end\{document\}', '', text)
    text = re.sub(r'\\section\*?\{([^}]*)\}', r'\n=== \1 ===\n', text)
    text = re.sub(r'\\subsection\*?\{([^}]*)\}', r'\n--- \1 ---\n', text)
    text = re.sub(r'\\textbf\{([^}]*)\}', r'**\1**', text)
    text = re.sub(r'\\textit\{([^}]*)\}', r'*\1*', text)
    text = re.sub(r'\\begin\{itemize\}', '', text)
    text = re.sub(r'\\end\{itemize\}', '', text)
    text = re.sub(r'\\item\s*', '\n- ', text)
    text = re.sub(r'\\[a-zA-Z]+\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\[a-zA-Z]+', '', text)
    text = re.sub(r'[{}]', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def parse_resume_text(text: str) -> dict:
    lines = [l.strip() for l in text.split('\n') if l.strip()]

    name = ""
    email = ""
    phone = ""
    location = ""
    linkedin = ""
    website = ""

    email_match = re.search(r'[\w.-]+@[\w.-]+\.\w+', text)
    if email_match:
        email = email_match.group(0)

    phone_match = re.search(r'[\+]?[\d\-\(\)\s]{7,15}', text)
    if phone_match:
        phone = phone_match.group(0).strip()

    linkedin_match = re.search(r'linkedin\.com/in/[\w-]+', text, re.I)
    if linkedin_match:
        linkedin = "https://" + linkedin_match.group(0)

    github_match = re.search(r'github\.com/[\w-]+', text, re.I)
    if github_match:
        website = "https://" + github_match.group(0)

    for line in lines[:5]:
        if re.search(r'@|\.com|linkedin|github|\d{3}', line):
            continue
        if len(line) > 2 and len(line) < 60 and not line.startswith(('-', '•', '*', '|')):
            name = line.replace('**', '').replace('*', '')
            break

    location_match = re.search(r'(?:Location|Address|City)[:\s]*(.+)', text, re.I)
    if location_match:
        location = location_match.group(1).strip()

    sections = {}
    current_section = "header"
    current_content = []

    section_patterns = [
        r'experience', r'work\s+history', r'employment', r'professional\s+experience',
        r'education', r'academic', r'qualification',
        r'skills', r'technical\s+skills', r'competencies', r'technologies',
        r'projects', r'portfolio', r'personal\s+projects',
        r'certifications?', r'licenses?', r'credentials?',
        r'summary', r'objective', r'profile', r'about',
    ]

    for line in lines:
        line_lower = line.lower().replace('*', '').replace('#', '').strip()
        matched_section = None
        for pattern in section_patterns:
            if re.match(rf'^{pattern}[:\s]*$', line_lower) or re.match(rf'^{pattern}$', line_lower):
                matched_section = pattern.split('\\s+')[0]
                break

        if matched_section:
            if current_content:
                sections[current_section] = '\n'.join(current_content)
            current_section = matched_section
            current_content = []
        else:
            current_content.append(line)

    if current_content:
        sections[current_section] = '\n'.join(current_content)

    summary = ""
    for key in ['summary', 'objective', 'profile', 'about']:
        if key in sections:
            summary = sections[key].strip()
            break

    experience = []
    exp_text = ""
    for key in ['experience', 'work', 'employment', 'professional']:
        if key in sections:
            exp_text = sections[key]
            break
    if exp_text:
        blocks = re.split(r'\n(?=\S.*(?:\d{4}|Present|Current|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec))', exp_text)
        for block in blocks:
            block_lines = [l.strip().lstrip('-•* ') for l in block.split('\n') if l.strip()]
            if not block_lines:
                continue
            title_line = block_lines[0]
            company = ""
            position = title_line
            date_match = re.search(r'(\w+\s+\d{4}\s*[-–]\s*(?:Present|Current|\w+\s+\d{4}))', title_line)
            if date_match:
                position = title_line[:date_match.start()].strip().rstrip(',|')
            parts = re.split(r'[\|@,]', position, maxsplit=1)
            if len(parts) == 2:
                position = parts[0].strip()
                company = parts[1].strip()
            elif len(parts) == 1 and ' at ' in position.lower():
                parts = position.split(' at ', 1)
                position = parts[0].strip()
                company = parts[1].strip()
            dates = ""
            date_match2 = re.search(r'(\w+\s+\d{4}\s*[-–]\s*(?:Present|Current|\w+\s+\d{4}))', block)
            if date_match2:
                dates = date_match2.group(1)
            desc_lines = block_lines[1:]
            description = "\n".join(desc_lines)
            highlights = [l for l in desc_lines if l.startswith(('- ', '• '))]
            experience.append({
                "company": company,
                "position": position,
                "startDate": dates.split('-')[0].strip() if '-' in dates or '–' in dates else "",
                "endDate": dates.split('-')[-1].strip() if '-' in dates or '–' in dates else "",
                "description": description,
                "highlights": highlights,
            })

    education = []
    edu_text = ""
    for key in ['education', 'academic', 'qualification']:
        if key in sections:
            edu_text = sections[key]
            break
    if edu_text:
        blocks = re.split(r'\n(?=\S)', edu_text)
        for block in blocks:
            block_lines = [l.strip().lstrip('-•* ') for l in block.split('\n') if l.strip()]
            if not block_lines:
                continue
            first_line = block_lines[0]
            institution = first_line
            degree = ""
            field = ""
            degree_match = re.search(r'(Bachelor|Master|PhD|B\.?S\.?|M\.?S\.?|B\.?A\.?|M\.?A\.?|Associate|Diploma)[^,]*', first_line, re.I)
            if degree_match:
                degree = degree_match.group(0).strip()
                parts = first_line.split(',', 1)
                if len(parts) == 2:
                    institution = parts[0].strip()
                    rest = parts[1].strip()
                    if degree_match.start() > 0:
                        field = first_line[degree_match.start():].split(',')[0].strip()
            education.append({
                "institution": institution,
                "degree": degree,
                "field": field,
                "startDate": "",
                "endDate": "",
                "gpa": "",
            })

    skills = []
    skills_text = ""
    for key in ['skills', 'technical', 'competencies', 'technologies']:
        if key in sections:
            skills_text = sections[key]
            break
    if skills_text:
        items = re.split(r'[,|\n]', skills_text)
        for item in items:
            skill = item.strip().lstrip('-•* ').strip()
            if skill and len(skill) > 1 and len(skill) < 50:
                category = ""
                if any(w in skill.lower() for w in ['python', 'java', 'javascript', 'typescript', 'go', 'rust', 'c++', 'ruby']):
                    category = "language"
                elif any(w in skill.lower() for w in ['react', 'vue', 'angular', 'django', 'flask', 'express', 'spring']):
                    category = "framework"
                elif any(w in skill.lower() for w in ['sql', 'mysql', 'postgres', 'mongo', 'redis', 'dynamodb']):
                    category = "database"
                elif any(w in skill.lower() for w in ['docker', 'kubernetes', 'aws', 'gcp', 'azure', 'jenkins', 'git']):
                    category = "tool"
                skills.append({"name": skill, "proficiency": "intermediate", "category": category})

    projects = []
    proj_text = ""
    for key in ['projects', 'portfolio', 'personal']:
        if key in sections:
            proj_text = sections[key]
            break
    if proj_text:
        blocks = re.split(r'\n(?=\S.*(?:\||–|-))', proj_text)
        if len(blocks) <= 1:
            blocks = re.split(r'\n(?=\S)', proj_text)
        for block in blocks:
            block_lines = [l.strip().lstrip('-•* ') for l in block.split('\n') if l.strip()]
            if not block_lines:
                continue
            first_line = block_lines[0]
            proj_name = first_line.split('|')[0].split('–')[0].split('-')[0].strip()
            desc_lines = block_lines[1:]
            projects.append({
                "name": proj_name,
                "description": "\n".join(desc_lines),
                "technologies": "",
                "link": "",
            })

    return {
        "title": f"{name} Resume" if name else "Uploaded Resume",
        "template": "modern",
        "summary": summary,
        "target_role": "",
        "personal": {
            "name": name,
            "email": email,
            "phone": phone,
            "location": location,
            "website": website,
            "linkedin": linkedin,
        },
        "experience": experience,
        "education": education,
        "skills": skills,
        "certifications": [],
        "projects": projects,
    }


@router.post("/", response_model=ResumeResponse, status_code=status.HTTP_201_CREATED)
async def create_resume(
    resume_data: ResumeCreate,
    current_user_id: uuid.UUID = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    return await resume_service.create_resume(db, current_user_id, resume_data)


@router.get("/", response_model=list[ResumeListResponse])
async def list_resumes(
    current_user_id: uuid.UUID = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    return await resume_service.list_user_resumes(db, current_user_id)


@router.get("/{resume_id}", response_model=ResumeResponse)
async def get_resume(
    resume_id: uuid.UUID,
    current_user_id: uuid.UUID = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    resume = await resume_service.get_resume(db, resume_id, current_user_id)
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    return resume


@router.put("/{resume_id}", response_model=ResumeResponse)
async def update_resume(
    resume_id: uuid.UUID,
    resume_data: ResumeUpdate,
    current_user_id: uuid.UUID = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    resume = await resume_service.get_resume(db, resume_id, current_user_id)
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    return await resume_service.update_resume(db, resume, resume_data)


@router.delete("/{resume_id}", response_model=ResumeResponse)
async def delete_resume(
    resume_id: uuid.UUID,
    current_user_id: uuid.UUID = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    resume = await resume_service.get_resume(db, resume_id, current_user_id)
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    return await resume_service.delete_resume(db, resume)


@router.post("/{resume_id}/experiences", response_model=ExperienceResponse, status_code=status.HTTP_201_CREATED)
async def add_experience(
    resume_id: uuid.UUID,
    data: ExperienceCreate,
    current_user_id: uuid.UUID = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    resume = await resume_service.get_resume(db, resume_id, current_user_id)
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    return await resume_service.add_experience(db, resume_id, data)


@router.post("/{resume_id}/educations", response_model=EducationResponse, status_code=status.HTTP_201_CREATED)
async def add_education(
    resume_id: uuid.UUID,
    data: EducationCreate,
    current_user_id: uuid.UUID = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    resume = await resume_service.get_resume(db, resume_id, current_user_id)
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    return await resume_service.add_education(db, resume_id, data)


@router.post("/{resume_id}/skills", response_model=SkillResponse, status_code=status.HTTP_201_CREATED)
async def add_skill(
    resume_id: uuid.UUID,
    data: SkillCreate,
    current_user_id: uuid.UUID = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    resume = await resume_service.get_resume(db, resume_id, current_user_id)
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    return await resume_service.add_skill(db, resume_id, data)


@router.post("/{resume_id}/certifications", response_model=CertificationResponse, status_code=status.HTTP_201_CREATED)
async def add_certification(
    resume_id: uuid.UUID,
    data: CertificationCreate,
    current_user_id: uuid.UUID = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    resume = await resume_service.get_resume(db, resume_id, current_user_id)
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    return await resume_service.add_certification(db, resume_id, data)


@router.post("/{resume_id}/projects", response_model=ProjectResponse, status_code=status.HTTP_201_CREATED)
async def add_project(
    resume_id: uuid.UUID,
    data: ProjectCreate,
    current_user_id: uuid.UUID = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    resume = await resume_service.get_resume(db, resume_id, current_user_id)
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    return await resume_service.add_project(db, resume_id, data)


@router.get("/{resume_id}/export")
async def export_resume(
    resume_id: uuid.UUID,
    current_user_id: uuid.UUID = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    resume = await resume_service.get_resume(db, resume_id, current_user_id)
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    content = resume.content or {}
    personal = content.get("personal", {}) if isinstance(content, dict) else {}

    export_data = {
        "title": resume.title,
        "template": resume.template,
        "summary": resume.summary or "",
        "target_role": resume.target_role or "",
        "personal": {
            "name": personal.get("name", ""),
            "email": personal.get("email", ""),
            "phone": personal.get("phone", ""),
            "location": personal.get("location", ""),
            "website": personal.get("website", ""),
            "linkedin": personal.get("linkedin", ""),
        },
        "experience": [
            {
                "company": e.company,
                "position": e.position,
                "startDate": e.start_date or "",
                "endDate": e.end_date or "",
                "description": e.description or "",
                "highlights": e.highlights or [],
            }
            for e in (resume.experiences or [])
        ],
        "education": [
            {
                "institution": e.institution,
                "degree": e.degree or "",
                "field": e.field_of_study or "",
                "startDate": e.start_date or "",
                "endDate": e.end_date or "",
                "gpa": e.gpa or "",
            }
            for e in (resume.educations or [])
        ],
        "skills": [
            {"name": s.name, "proficiency": s.proficiency or "intermediate", "category": s.category or ""}
            for s in (resume.skills or [])
        ],
        "certifications": [
            {"name": c.name, "issuer": c.issuer or ""}
            for c in (resume.certifications or [])
        ],
        "projects": [
            {
                "name": p.name,
                "description": p.description or "",
                "technologies": ", ".join(str(t) for t in (p.technologies or [])) if isinstance(p.technologies, list) else (p.technologies or ""),
                "link": p.live_url or p.github_url or "",
            }
            for p in (resume.projects or [])
        ],
    }

    filename = f"{(resume.title or 'resume').replace(' ', '_').lower()}.json"
    content_bytes = json.dumps(export_data, indent=2).encode("utf-8")
    return StreamingResponse(
        BytesIO(content_bytes),
        media_type="application/json",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/import", response_model=ResumeResponse, status_code=status.HTTP_201_CREATED)
async def import_resume(
    file: UploadFile = File(...),
    current_user_id: uuid.UUID = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    filename = file.filename or ""
    raw = await file.read()

    if len(raw) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large. Maximum size is 10MB.")

    try:
        if filename.endswith(".json"):
            try:
                data = json.loads(raw.decode("utf-8"))
            except (json.JSONDecodeError, UnicodeDecodeError):
                raise HTTPException(status_code=400, detail="Invalid JSON file")
        elif filename.endswith(".pdf"):
            text = extract_text_from_pdf(raw)
            if text is None:
                raise HTTPException(status_code=400, detail="Failed to parse PDF. Ensure the file is a valid PDF and PyMuPDF is installed.")
            if not text.strip():
                raise HTTPException(status_code=400, detail="PDF appears to be empty or contains only images.")
            print(f"Attempting LLM-assisted parsing for PDF ({len(text)} chars)...")
            data = llm_parse_resume_text(text)
            if data:
                print("LLM parsing succeeded")
            else:
                print("LLM parsing failed or unavailable, falling back to regex")
                data = parse_resume_text(text)
        elif filename.endswith(".docx"):
            text = extract_text_from_docx(raw)
            if text is None:
                raise HTTPException(status_code=400, detail="Failed to parse DOCX. Ensure python-docx is installed.")
            if not text.strip():
                raise HTTPException(status_code=400, detail="DOCX file appears to be empty.")
            print(f"Attempting LLM-assisted parsing for DOCX ({len(text)} chars)...")
            data = llm_parse_resume_text(text)
            if data:
                print("LLM parsing succeeded")
            else:
                print("LLM parsing failed or unavailable, falling back to regex")
                data = parse_resume_text(text)
        elif filename.endswith(".tex") or filename.endswith(".latex"):
            try:
                content = raw.decode("utf-8")
            except UnicodeDecodeError:
                raise HTTPException(status_code=400, detail="Failed to decode TeX file")
            text = extract_text_from_tex(content)
            print(f"Attempting LLM-assisted parsing for TeX ({len(text)} chars)...")
            data = llm_parse_resume_text(text)
            if data:
                print("LLM parsing succeeded")
            else:
                print("LLM parsing failed or unavailable, falling back to regex")
                data = parse_resume_text(text)
        else:
            raise HTTPException(status_code=400, detail="Supported formats: JSON, PDF, DOCX, TeX")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse file: {str(e)}")

    try:
        personal = data.get("personal", {}) or {}
        personal = {k: v for k, v in personal.items() if v}

        resume_data = ResumeCreate(
            title=data.get("title") or (filename.rsplit('.', 1)[0].replace('_', ' ').title() if filename else "Uploaded Resume"),
            template=data.get("template", "modern"),
            summary=data.get("summary", ""),
            target_role=data.get("target_role", ""),
            content={"personal": personal} if personal else {"personal": {}},
        )

        resume = await resume_service.create_resume(db, current_user_id, resume_data)

        for exp in (data.get("experience") or data.get("experiences") or []):
            try:
                if isinstance(exp, dict) and (exp.get("company") or exp.get("position")):
                    await resume_service.add_experience(db, resume.id, ExperienceCreate(
                        company=exp.get("company", "") or "",
                        position=exp.get("position") or exp.get("title", "") or "",
                        start_date=exp.get("startDate") or exp.get("start_date", ""),
                        end_date=exp.get("endDate") or exp.get("end_date", ""),
                        description=exp.get("description", ""),
                        highlights=exp.get("highlights") or [],
                    ))
            except Exception as e:
                print(f"Failed to add experience item: {e}")

        for edu in (data.get("education") or data.get("educations") or []):
            try:
                if isinstance(edu, dict) and edu.get("institution"):
                    await resume_service.add_education(db, resume.id, EducationCreate(
                        institution=edu["institution"],
                        degree=edu.get("degree", ""),
                        field_of_study=edu.get("field") or edu.get("field_of_study", ""),
                        start_date=edu.get("startDate") or edu.get("start_date", ""),
                        end_date=edu.get("endDate") or edu.get("end_date", ""),
                        gpa=edu.get("gpa", ""),
                    ))
            except Exception as e:
                print(f"Failed to add education item: {e}")

        for skill in (data.get("skills") or []):
            try:
                if isinstance(skill, dict) and skill.get("name"):
                    await resume_service.add_skill(db, resume.id, SkillCreate(
                        name=skill["name"],
                        proficiency=skill.get("proficiency", "intermediate"),
                        category=skill.get("category", ""),
                    ))
                elif isinstance(skill, str) and skill.strip():
                    await resume_service.add_skill(db, resume.id, SkillCreate(name=skill.strip()))
            except Exception as e:
                print(f"Failed to add skill item: {e}")

        for cert in (data.get("certifications") or []):
            try:
                if isinstance(cert, dict) and cert.get("name"):
                    await resume_service.add_certification(db, resume.id, CertificationCreate(
                        name=cert["name"],
                        issuer=cert.get("issuer", ""),
                    ))
            except Exception as e:
                print(f"Failed to add certification item: {e}")

        for proj in (data.get("projects") or []):
            try:
                if isinstance(proj, dict) and proj.get("name"):
                    techs = proj.get("technologies", "")
                    if isinstance(techs, str):
                        techs = [t.strip() for t in techs.split(",") if t.strip()]
                    await resume_service.add_project(db, resume.id, ProjectCreate(
                        name=proj["name"],
                        description=proj.get("description", ""),
                        technologies=techs if isinstance(techs, list) else [],
                        github_url=proj.get("github_url", ""),
                        live_url=proj.get("link") or proj.get("live_url", ""),
                    ))
            except Exception as e:
                print(f"Failed to add project item: {e}")

        await db.commit()
        return await resume_service.get_resume(db, resume.id, current_user_id)
    except HTTPException:
        raise
    except Exception as e:
        await db.rollback()
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to save imported resume: {str(e)}")


from pydantic import BaseModel


class AIEnhanceRequest(BaseModel):
    text: str
    context: str = "summary"
    skills: list[str] = []
    project_name: str = ""
    project_tech: str = ""


@router.post("/ai-enhance")
async def ai_enhance_text(
    request: AIEnhanceRequest,
    current_user_id: uuid.UUID = Depends(get_current_user),
):
    try:
        import ollama as ollama_sync
        from app.core.config import settings

        if request.context == "summary":
            skills_str = ", ".join(request.skills) if request.skills else "various technologies"
            prompt = f"""You are an expert resume writer. Enhance the following professional summary to be concise, impactful, and ATS-friendly.

Current summary: {request.text or '(empty - write a new one)'}
Skills: {skills_str}

Rules:
- 2-3 sentences maximum
- Start with a strong action-oriented statement
- Include key technical skills naturally
- Use professional language
- No bullet points, just flowing text
- Do NOT use phrases like "I am" or "I have" - use第三人称 or implied subject
- Return ONLY the enhanced summary text, nothing else"""

        elif request.context == "project":
            prompt = f"""You are an expert resume writer. Write a compelling 2-line project description for a resume.

Project name: {request.project_name}
Technologies: {request.project_tech}
Current description: {request.text or '(empty)'}

Rules:
- Exactly 2 sentences
- First sentence: what the project does and its impact
- Second sentence: key technical details and achievements
- Use action verbs (Built, Developed, Engineered, Implemented, Designed)
- Include quantifiable results where possible
- Professional tone, no filler words
- Return ONLY the description text, nothing else"""

        else:
            return {"enhanced_text": request.text}

        model = settings.DEFAULT_MODEL
        response = ollama_sync.chat(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            options={"num_gpu": 99, "temperature": 0.7}
        )
        enhanced = response["message"]["content"].strip()

        enhanced = enhanced.strip('"').strip("'")
        if enhanced.startswith("```"):
            enhanced = enhanced.split("\n", 1)[-1].rsplit("```", 1)[0].strip()

        return {"enhanced_text": enhanced}
    except ImportError:
        raise HTTPException(status_code=500, detail="Ollama not installed. AI enhancement unavailable.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI enhancement failed: {str(e)}")
